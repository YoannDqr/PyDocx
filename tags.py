from docx.shared import Cm
from docx.table import _Cell
from tools import *


def title(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        return None
    current = document.add_heading(value, int(option['rank']))
    return current


def img(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        return None
    if not('path' in option):
        return None
    try:
        border = int(option['border'])
    except:
        border = 0
    try:
        width = Cm(int(option['width']))
    except:
        width = None
    caption = not('caption' in option)

    add_picture(
        document,
        path=option['path'],
        caption=value,
        width=width,
        border=border,
        caption_switch=caption
    )


def styled_string(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        return None

    added = document.paragraphs[-1].add_run(value)
    apply_legacy_style(added, parents_node)


def p(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        if type(document) != _Cell or document.paragraphs[-1].text != "":
            document.add_paragraph('')

    else:
        document.paragraphs[-1].add_run(value)


def testssl(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        file = open(option['path'], 'r')
        file_content = file.readlines()
        file.close()

        flag_cipher_suites = False
        flag_key_size = False
        cipher_suites = {}
        cipher_keys = {}
        protocols = []
        for elt in file_content:
            if 'Cipher order' in elt:
                flag_cipher_suites = True
            elif 'Testing server defaults' in elt:
                flag_cipher_suites = False
            elif '-----------------' in elt:
                flag_key_size = True
            elif 'Running' in elt:
                break

            elif flag_key_size and elt.strip() != '':
                splitted = [var.strip() for var in elt.split(' ') if var.strip() != '']
                cipher_suite = splitted[1]
                key = splitted[-2]
                cipher_keys[cipher_suite] = key

            elif flag_cipher_suites and elt.strip() != '':
                if ':' in elt:
                    splitted = elt.split(':')
                    protocol = splitted[0].strip()
                    protocols.append(protocol)
                    cipher_list = [val.strip() for val in splitted[1].split(' ') if val.strip() != '']
                else:
                    cipher_list = [val.strip() for val in elt.strip().split(' ') if val.strip() != '']

                for val in cipher_list:
                    try:
                        cipher_suites[val] += [protocol]
                    except:
                        cipher_suites[val] = [protocol]

        table = document.add_table(rows=len(cipher_suites)+1, cols=len(protocols)+2)
        table.cell(0, 0).paragraphs[-1].add_run('Suite de chiffrement')
        table.cell(0, 1).paragraphs[-1].add_run('Taille de clé')

        for i in range(len(protocols)):
            table.cell(0, i + 2).paragraphs[-1].add_run(protocols[i])

        i = 0
        for key, value in cipher_suites.items():
            table.cell(i + 1, 0).paragraphs[-1].add_run(key)
            table.cell(i + 1, 1).paragraphs[-1].add_run(cipher_keys[key])
            for elt in value:
                table.cell(i + 1, 2 + protocols.index(elt)).paragraphs[-1].add_run('X')
            i += 1

    else:
        if value.strip() != '':
            add_caption(document, 'Table', value)
        else:
            document.add_paragraph('')


def nmap(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        file = open(option['path'], 'r')
        file_content = file.readlines()
        file.close()
        read = False
        data = []
        for elt in file_content:
            if 'PORT' in elt and 'STATE' in elt and 'SERVICE' in elt:
                read = True
            elif elt == '\n':
                read = False
            elif read:
                data.append([val.strip() for val in elt.split(' ') if val != ''])
        table = document.add_table(rows=len(data)+1, cols=len(data[0])+2)
        table.rows[0].cells[-1].paragraphs[-1].add_run('Commentaire')
        data = [['Protocol', 'Port', 'Etat', 'Service']] + data
        for i in range(len(data)):
            try:
                port, protocol = data[i][0].split('/')
                data[i][0] = port
                data[i] = [protocol]+data[i]
            except ValueError:
                pass
            for j in range(len(data[0])):
                table.cell(i, j).paragraphs[-1].add_run(data[i][j])

    else:
        if value.strip() != '':
            add_caption(document, 'Table', value)
        else:
            document.add_paragraph('')


def table(document, option, value, parents_node, preprocessing=False):
    if preprocessing:
        if ('c' in option) and ('l' in option):
            document.add_table(rows=int(option['l']), cols=int(option['c']))
            if 'caption' in option:
                add_caption(document, 'Table', option['caption'])
            option['c'], option['l'] = 0, 0
    else:
        split_cells = option['delimiter']
        elt = value.strip()
        if elt == split_cells:
            option['c'] += 1

        if option['c'] == len(document.tables[-1].rows[0].cells) - 1 and not (elt == split_cells):
            option['l'] += 1
            option['c'] = 0


def stall(document, option, value, parents_node, preprocessing=False):
    pass


