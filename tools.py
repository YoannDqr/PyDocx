from docx.oxml import OxmlElement, ns
from PIL import Image, ImageOps
from docx.enum.style import WD_STYLE_TYPE


def get_cells_coordinate(table, area):
    # Return the raw and columns indexes given by the style string used by tables !11-15/*-12-15#...
    raw_rows, raw_columns = area.split('/')

    if '*' in raw_rows:  # All except those explicitly given
        rows = [i for i in range(len(table.rows)) if not (str(i) in raw_rows.split('-'))]
    else:
        rows = [int(row) for row in raw_rows.split('-')]
    if '*' in raw_columns:
        columns = [i for i in range(len(table.columns)) if not (str(i) in raw_columns.split('-'))]
    else:
        columns = [int(column) for column in raw_columns.split('-')]

    return rows, columns


def get_subdocument(document, parents_node):
    doc = document
    for elt in parents_node:
        if ('l' in elt) and ('c' in elt):
            a = doc.tables[-1].rows[elt['l']].cells
            doc = doc.tables[-1].rows[elt['l']].cells[elt['c']]
    return doc


def add_border(input_image, output_image, border):
    img = Image.open(input_image)
    if isinstance(border, int) or isinstance(border, tuple):
        bimg = ImageOps.expand(img, border=border, fill='black')
    else:
        raise RuntimeError('Border is not an image or tuple')
    bimg.save(output_image)


def add_xref(paragraph, type):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ '+type+' \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'end')
    r.append(fldChar)


def add_caption(document, type, caption):
    paragraph = document.add_paragraph(type, style='Caption')
    add_xref(paragraph, type)
    paragraph.add_run(': ' + caption)
    return paragraph


def add_picture(document, path="", caption="", width=None, border=0, caption_switch=True):
    tmp_path = path
    if border > 0:
        splitted = tmp_path.split('/')
        final_path = '/'.join(splitted[:-1]) + '/tmp_' + splitted[-1]
        add_border(tmp_path, final_path, border)
        tmp_path = final_path

    if len(document.paragraphs) == 0 or document.paragraphs[-1].text != '':
        run = document.add_paragraph('').add_run()
    else:
        run = document.paragraphs[-1].add_run()
    run.add_picture(tmp_path, width=width)
    if caption_switch:
        add_caption(document, 'Figure', caption)
    return None


def apply_legacy_style(elt, parents_node):
    for opt in parents_node:
        for key, value in opt.items():
            if 'run' in key:
                try:
                    getattr(elt, key.split('run_')[1])
                    setattr(elt, key.split('run_')[1], int(value))
                except AttributeError:
                    pass
                except ValueError:
                    setattr(elt, key.split('run_')[1], value)


def create_paragraph_style(paragraph):
    style = paragraph.style
    if style.builtin:
        i = 0
        success = False
        while success is not True:
            try:
                new_style = paragraph.part.document.styles.add_style('tmp_style{}'.format(i), WD_STYLE_TYPE.PARAGRAPH)
                success = True
            except:
                i += 1
        new_style.base_style = style
    else:
        new_style = style
    return new_style


def progress_bar(current, total, title="", barLength=20):
    percent = float(current) * 100 / total
    arrow = '-' * int(percent / 100 * barLength - 1) + '>'
    spaces = ' ' * (barLength - len(arrow))

    print(title + ' Progress: [%s%s] %d %%' % (arrow, spaces, percent), end='\r')

