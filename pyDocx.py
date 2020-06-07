from docx import Document
from settings import *
from tools import get_subdocument, progress_bar
import re
import multiprocessing
from docxcompose.composer import Composer
import os
import time
import sys


def generate_graph(line, parent=None):
    def process_option(option):
        option_args = []
        raw_classes = []
        result = (option_args, raw_classes)
        key = True
        index = 0
        tag_name = ""
        first = True
        if SPLIT_CLASS in option or SPLIT_OPTION in option:
            for i in range(len(option)):
                if option[i] == SPLIT_CLASS or option[i] == SPLIT_OPTION:
                    if first:
                        tag_name = option[:i]
                        first = False
                    else:
                        result[key].append(option[index:i])
                    index = i + 1
                    key = option[i] == SPLIT_CLASS

            result[key].append(option[index:])
        else:
            tag_name = option

        options = {'name': tag_name}
        styles = {}
        classes = []

        for elt in option_args:
            splitted = elt.split(SPLIT_OPTION_VALUE)
            if len(splitted) > 1:
                options[splitted[0]] = " ".join(splitted[1].split(SPACE_CHAR))

        for elt in raw_classes:
            splitted = elt.split(SPLIT_STYLE_VALUE)
            if len(splitted) == 2:
                styles[splitted[0]] = " ".join(splitted[1].split(SPACE_CHAR))
            else:
                classes.append(elt)

        return options, classes, styles

    if parent is None:
        parent = []
    result = []
    index = line.find(TAG_CHAR)
    while index != -1:
        option_end = line.find(' ', index)
        tmp = line.find('\n', index)
        if tmp < option_end and tmp != -1:
            option_end = tmp

        option, classes, styles = process_option(line[index + len(TAG_CHAR): option_end])

        end_index = line.find(TAG_CHAR_END + option['name'], index + 1)
        subline = line[option_end + 1:end_index - 1]
        while subline.count(TAG_CHAR + option['name']) != subline.count(TAG_CHAR_END + option['name']):
            end_index = line.find(TAG_CHAR_END + option['name'], index + 1 + end_index)
            subline = line[option_end + 1:end_index - 1]

        new_parent = parent + [option]
        node_header = {'option': option, 'class': classes, 'style': styles}
        if index == 0 or parent == []:
            result += [[node_header] + generate_graph(subline, parent=new_parent)]
        else:
            result += [line[:index]] + [
                [node_header] + generate_graph(subline, parent=new_parent)]

        if end_index + len(TAG_CHAR) + len(option['name']) < len(line):
            line = line[end_index + len(TAG_CHAR) + len(option['name']):]
        else:
            line = ''
        index = line.find(TAG_CHAR)

    if line != '':
        result += [line]
    return result


def graph2doc(document, graph, current, gstyle=None, parents_node=None):
    def apply_global_style(document, classe, gstyle, inline_style):
        option_name = classe[-1]
        for elt in classe:
            if elt in gstyle:
                for key, value in gstyle[elt].items():
                    if not (key in inline_style):
                        try:
                            STYLE_FUNCTIONS[option_name][key](document, value)
                        except KeyError:
                            print("[W] - Style property `{}={}` is not implemented for element `{}`".format(key, value,
                                                                                                            option_name))

        for key, value in inline_style.items():
            try:
                STYLE_FUNCTIONS[option_name][key](document, value)
            except KeyError:
                print("[W] - Style property `{}={}` is not implemented for element `{}`".format(key, value, option_name))

    if parents_node is None:
        parents_node = []
    if gstyle is None:
        gstyle = {}

    option = graph[0]['option']
    classes = graph[0]['class']
    styles = graph[0]['style']

    subdoc = get_subdocument(document, parents_node)
    try:
        TAGS_FUNCTIONS[option['name']](subdoc, option, None, parents_node, preprocessing=True)
    except KeyError:
        print('The tag {} is not recognized.'.format(option['name']))
        exit()
    for elt in graph[1:]:
        # If the function reach a leave
        if type(elt) == str:
            try:
                TAGS_FUNCTIONS[option['name']](subdoc, option, elt, parents_node + [option])
            except KeyError:
                print('The tag {} is not recognized.'.format(option['name']))
                exit()
            apply_global_style(subdoc, classes + [option['name']], gstyle, styles)
        elif type(elt) == list:
            apply_global_style(subdoc, classes + [option['name']], gstyle, styles)
            # Recursively parse sons trees
            graph2doc(document, elt, current, gstyle=gstyle, parents_node=parents_node + [option])


def initial_parsing(initial_data):
    def process_aliases(data):
        for key, value in ALIAS.items():

            pattern = re.compile(rf'({re.escape(TAG_CHAR)}){key}([\?:\s\n])')
            data = pattern.sub(rf'\1{value[0]}\2', data)

            pattern = re.compile(rf'({re.escape(TAG_CHAR_END)}){key}([\?:\s\n])')
            data = pattern.sub(rf'\1{value[1]}\2', data)

        return data

    data = ""
    style = {}
    keep_style = False
    for elt in initial_data:
        if elt != '\n':
            if TAG_CHAR + 'style' in elt:
                keep_style = True
            elif TAG_CHAR_END + 'style' in elt:
                keep_style = False
            elif keep_style:
                key, value = elt.split(SPLIT_STYLE)
                for style_opt in value.strip().split(SPLIT_STYLE_PROPS):
                    prop, val = style_opt.split(SPLIT_STYLE_VALUE)
                    try:
                        style[key.strip()][prop] = val
                    except:
                        style[key.strip()] = {prop: val}
            else:
                data += elt.lstrip()

    data = process_aliases(data)

    return style, data + TAG_CHAR + "p " + TAG_CHAR_END + "p"


def task(id, graph, gstyle):
    doc = Document('docx/template.docx')
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)
    p._p = p._element = None
    graph2doc(doc, graph, None, gstyle=gstyle)
    doc.save('docx/tmp/part{}.docx'.format(id))


def gen_doc(graph_data, style, template=None, max_processes=os.cpu_count()+1):
    if template is None:
        doc = Document()
    else:
        doc = Document(template)

    composer = Composer(doc)
    workers = []
    progress = 0
    for i in range(len(graph_data)):
        process = multiprocessing.Process(target=task, args=(i, graph_data[i], style))
        workers.append(process)
        process.start()
        if i == max_processes:
            for elt in workers:
                elt.join()
                composer.append(Document('docx/tmp/part{}.docx'.format(progress)))
                progress += 1
            workers = []
    for elt in workers:
        elt.join()
        composer.append(Document('docx/tmp/part{}.docx'.format(progress)))
        progress += 1

    composer.save('docx/final.docx')
    print("Document generated in docx/final.docx.\n")


def translate_document(path, template=None, encoding='utf-8'):
    file = open(path, encoding=encoding)
    style, data = initial_parsing(file.readlines())
    file.close()
    graph_data = generate_graph(data)
    gen_doc(graph_data, style, template)


if __name__ == '__main__':
    if len(sys.argv) == 2:
        template = None
        file = sys.argv[1]
    elif len(sys.argv) == 3:
        file = sys.argv[1]
        template = sys.argv[2]
    else:
        file = 'template.dqr'
        template = 'docx/template.docx'

    if os.path.isfile(file) is not True:
        print("The file {} does not exist...".format(file))
        exit()
    if os.path.isfile(template) is not True:
        print("The file {} does not exist...".format(template))
        exit()
    if os.path.isdir('docx/tmp') is not True:
        os.mkdir('docx/tmp')

    start_time = time.time()
    translate_document(file, template)
    for elt in os.scandir('docx/tmp'):
        os.remove(elt)
    print("Multiprocess runtime : %s seconds" % (time.time() - start_time))


